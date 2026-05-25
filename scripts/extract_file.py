import base64
import io
import json
import sys
import zipfile
import xml.etree.ElementTree as ET


def extract_docx(blob):
    from docx import Document

    document = Document(io.BytesIO(blob))
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(document.tables, 1):
        parts.append(f"\n[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx(blob):
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(blob), read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"\n[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                parts.append(" | ".join(values).rstrip(" |"))
    return "\n".join(parts)


def extract_pptx(blob):
    namespaces = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    parts = []
    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        slide_names = sorted(
            name for name in archive.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        for index, name in enumerate(slide_names, 1):
            root = ET.fromstring(archive.read(name))
            texts = []
            for node in root.findall(".//a:t", namespaces):
                if node.text and node.text.strip():
                    texts.append(node.text.strip())
            if texts:
                parts.append(f"\n[Slide {index}]\n" + "\n".join(texts))
    return "\n".join(parts)


def main():
    payload = json.loads(sys.stdin.read())
    filename = payload.get("filename", "")
    extension = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    blob = base64.b64decode(payload.get("base64", ""))

    if extension == "docx":
        text = extract_docx(blob)
    elif extension == "xlsx":
        text = extract_xlsx(blob)
    elif extension == "pptx":
        text = extract_pptx(blob)
    elif extension in {"doc", "xls", "ppt"}:
        raise ValueError("暂不支持旧版 Office 二进制格式，请另存为 .docx / .xlsx / .pptx 后上传。")
    else:
        text = blob.decode("utf-8", errors="ignore")

    print(json.dumps({"filename": filename, "text": text}, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as error:
        print(json.dumps({"error": str(error)}, ensure_ascii=False))
        sys.exit(1)
